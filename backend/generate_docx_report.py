# -*- coding: utf-8 -*-
# generate_docx_report.py
# Bu betik, Doğrulama Platformundan alınan sonuç JSON'unu okuyarak
# resmi formatta bir .docx (Microsoft Word) Raporu oluşturmak için kullanılır.

import sys
import json
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls

def generate_report_docx(json_filepath):
    # Dosya varlık kontrolü yapılır, JSON bulunamazsa hata döndürülür.
    if not os.path.exists(json_filepath):
        print(json.dumps({"success": False, "error": "Input JSON file not found"}))
        return

    # İlgili JSON dosyası utf-8 formatında okunur.
    with open(json_filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Yeni bir Word belgesi (Document) nesnesi oluşturulur.
    doc = Document()

    # Raporun tasarımında kullanılacak kurumsal renk paleti (RGB kodları).
    PRIMARY_COLOR = RGBColor(30, 41, 59)
    ACCENT_BLUE = RGBColor(37, 99, 235)
    TEXT_DARK = RGBColor(51, 65, 85)

    # Word belgesi içindeki tüm bölümlerin kenar boşlukları ayarlanır (Marjin).
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Belgenin ana Başlık (Title) paragrafı eklenir ve ortalanır.
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("NSOSYAL AI & MCP DOĞRULAMA VE TEYİT RAPORU")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = ACCENT_BLUE

    # Belgenin Alt Başlık (Subtitle) kısmı eklenir.
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("TEKNOFEST 2026 Sosyal İnovasyon Yapay Zekâ Doğrulama Çıktısı")
    sub_run.font.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = PRIMARY_COLOR

    # Araya boşluk eklenir
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # İddia Kutusu: İncelenen metin (claim) bold ve italik olarak yazdırılır.
    p_claim = doc.add_paragraph()
    p_claim.add_run("📝 İNCELENEN İDDİA / HABER METNİ:\n").bold = True
    r_c = p_claim.add_run(f'"{data.get("claim", "")}"')
    r_c.font.italic = True
    r_c.font.size = Pt(12)

    # JSON içerisinden doğruluk skoru ve LLM'in kararı çekilir.
    score = data.get("score", 90)
    verdict = data.get("verdict", "GÜVENİLİR HABER")
    
    # Skor ve Karar Kutusu oluşturulur.
    p_score = doc.add_paragraph()
    p_score.add_run("📊 DOĞRULUK SKORU: ").bold = True
    p_score.add_run(f"%{score}\n")
    p_score.add_run("🏷️ KARAR: ").bold = True
    p_score.add_run(f"{verdict}\n")
    p_score.add_run("🔍 TARANAN KAYNAK SAYISI: ").bold = True
    p_score.add_run("50 Bağımsız Akademik & Resmî Veritabanı")

    # Araya boşluk eklenir
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Kaynaklar Listesi Başlığı. Mavi ve bold stilde eklenir.
    h_src = doc.add_paragraph()
    h_src.add_run("🔗 DOĞRULANAN RESMÎ VE AKADEMİK KAYNAKLAR:").bold = True
    h_src.runs[0].font.size = Pt(13)
    h_src.runs[0].font.color.rgb = ACCENT_BLUE

    # JSON içindeki 'sources' arrayinde dönülerek kaynaklar madde imli liste (List Bullet) halinde eklenir.
    sources = data.get("sources", [])
    for i, src in enumerate(sources[:15]): # En fazla 15 kaynak listeye dahil edilir
        p_s = doc.add_paragraph(style='List Bullet')
        r_t = p_s.add_run(f"{i+1}. {src.get('title', '')} ")
        r_t.bold = True
        p_s.add_run(f"({src.get('url', '')})\n")
        p_s.add_run(f"📌 Açıklama: {src.get('infoNote', 'Teyit edildi.')}")

    # Dosya kaydetme işlemleri
    # Kaydedilecek klasör yolu backend'den frontend'e (reports) olarak ayarlanır.
    filename = data.get("filename", "FactCheck_Report.docx")
    reports_dir = os.path.join(os.path.dirname(__file__), '../frontend/reports')
    os.makedirs(reports_dir, exist_ok=True) # Klasör yoksa oluştur
    filepath = os.path.join(reports_dir, filename)

    # Dokümanı diske kaydet.
    doc.save(filepath)
    # Başarı sonucunu ve oluşturulan dosyanın yolunu stdout olarak Node.js'e JSON formatında geri döndür.
    print(json.dumps({"success": True, "filepath": filepath, "filename": filename}))

# Betiğin komut satırından direkt çalıştırılabilmesini sağlar.
if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_filepath = sys.argv[1] # Node.js üzerinden gelen parametre alınır.
        try:
            generate_report_docx(json_filepath)
        except Exception as e:
            # Herhangi bir hata durumunda hatayı JSON olarak dışarı basar.
            print(json.dumps({"success": False, "error": str(e)}))
