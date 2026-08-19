import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_team_table_docx(output_path):
    doc = docx.Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("TEKNOFEST 2026 - NSOSYAL İNOVASYON PROJESİ")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate Dark

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("TAKIM GÖREV DAĞILIMI VE DİSİPLİNLERARASI KATKI TABLOSU")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(13, 148, 136) # Teal Accent

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_note = p_note.add_run("Not: Değerlendirme esasları gereği takım üyelerinin isim, soyad ve fotoğraf gibi kişisel bilgilerine yer verilmeyerek anonim yapı korunmuştur.")
    run_note.font.name = "Calibri"
    run_note.font.size = Pt(10)
    run_note.font.italic = True
    run_note.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph() # Spacing

    # Table Setup
    table_data = [
        ("Sıra / Üye", "Takım Rolü", "Disiplin / Uzmanlık Alanı", "Projeye Doğrudan Katkısı"),
        (
            "1. Takım Üyesi",
            "Takım Kaptanı &\nYapay Zekâ Geliştiricisi\n(Backend)",
            "Yapay Zekâ, Agentic AI,\nRAG Mimarileri, Ürün Yönetimi",
            "Projenin genel mimari kurgusunu ve ürün vizyonunu yönetmiştir. Model Context Protocol (MCP v1.0) standartlarına uygun otonom ajan orkestrasyonunu, 50 kaynaklı doğrulama araçlarını, Tesseract Vision OCR ve LLM tabanlı multimodal doğrulama zincirini sıfırdan geliştirmeyi üstlenmiştir."
        ),
        (
            "2. Takım Üyesi",
            "Siber Güvenlik ve Altyapı Sorumlusu\n(Backend)",
            "Siber Güvenlik, Sistem Güvenliği,\nTehdit Analizi",
            "Nginx WAF ve Docker container altyapısının izolasyonunu ve güvenliğini sağlamıştır. IP tabanlı Rate Limiting (dakikada max 20 istek sınırı) DoS koruma katmanını kodlamış; Hugging Face URLBERT Transformer modelini Python Flask mikroservis olarak entegre ederek oltalama (phishing) engelleyici siber güvenlik katmanını kurmuştur."
        ),
        (
            "3. Takım Üyesi",
            "Veritabanı ve Veri Ön İşleme Sorumlusu\n(Backend)",
            "Veri Bilimi, Veritabanı Mimarisi,\nVeri Ön İşleme",
            "SQLite3 veritabanı şemasını, ilişkisel tabloları ve veritabanı önbellekleme (caching) mimarisini tasarlamıştır. 13 canlı kariyer platformundan periyodik veri çeken iş ilanı veritabanı boru hatlarını ve yapay zekâ moderasyonu için gerekli veri setlerini düzenlemiştir."
        ),
        (
            "4. Takım Üyesi",
            "Frontend Geliştiricisi\n(Frontend)",
            "Web Yazılım Geliştirme,\nUI/UX Tasarımı",
            "Kullanıcı dostu, modern ve duyarlı (responsive) web arayüz bileşenlerini tasarlamış ve kodlamıştır. Doğrulama sonuçlarının şeffaf biçimde sunulduğu skor kartları, kanıt detay pencereleri ve kullanıcı etkileşim bileşenlerini geliştirmeyi üstlenmiştir."
        ),
        (
            "5. Takım Üyesi",
            "Frontend Geliştiricisi\n(Frontend)",
            "Frontend Yazılım,\nKullanıcı Deneyimi (UX),\nAPI Entegrasyonu",
            "Uygulama içi tüm ekranların tasarımını ve istemci tarafı (client-side) REST API haberleşme entegrasyonlarını gerçekleştirmiştir. Görsel ve metin yükleme alanlarının kullanıcı deneyimini optimize etmiş, canlı doğrulama durumu göstergelerini ve hata kontrol ekranlarını kodlamıştır."
        )
    ]

    table = doc.add_table(rows=len(table_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Inches(1.2), Inches(1.8), Inches(1.8), Inches(3.2)]

    # Style Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(table_data[0]):
        hdr_cells[i].text = title
        shading = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Fill rows
    for r_idx in range(1, len(table_data)):
        row_cells = table.rows[r_idx].cells
        data = table_data[r_idx]

        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"

        for c_idx in range(4):
            row_cells[c_idx].text = data[c_idx]
            shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)

            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx < 3 else WD_ALIGN_PARAGRAPH.LEFT

            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                if c_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(15, 23, 42)
                else:
                    run.font.color.rgb = RGBColor(51, 65, 85)

    # Set column widths & cell margins
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m in ['top', 'bottom']:
                node = OxmlElement(f'w:{m}')
                node.set(qn('w:w'), '120')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            for m in ['left', 'right']:
                node = OxmlElement(f'w:{m}')
                node.set(qn('w:w'), '160')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

    doc.save(output_path)
    print("Word dosyasi basariyla olusturuldu.")

if __name__ == "__main__":
    create_team_table_docx("c:/Users/Huzur Bilgisayar/OneDrive/Masaüstü/site_devam/Takim_Gorev_Dagilimi_Tablosu.docx")
